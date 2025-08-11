#%%
import zipfile
import io
import pandas as pd
import requests
import re
import tempfile
from pathlib import Path
import warnings
warnings.filterwarnings('ignore')

BASE_DIR = Path(__file__).parent
TPL_PATH = BASE_DIR / "models" / "FICHA_FUNDO.docx"   # respeita maiúsculas/minúsculas

if not TPL_PATH.exists():
    st.error(f"Template não encontrado: {TPL_PATH}")
    st.stop()

TEMPLATE_BYTES = TPL_PATH.read_bytes()

#%%
url_data_registros = "https://dados.cvm.gov.br/dados/FI/CAD/DADOS/registro_fundo_classe.zip"


resp = requests.get(url_data_registros)


# %%

dfs = {}
with zipfile.ZipFile(io.BytesIO(resp.content)) as z:
    for name in z.namelist():
        if not name.lower().endswith(".csv"):
            continue
        with z.open(name) as f:
            text = f.read().decode("latin1", errors="replace")
        try:
            df = pd.read_csv(io.StringIO(text), sep=";", engine="python")
        except pd.errors.ParserError:
            # fallback para aspas/linhas corrompidas
            df = pd.read_csv(io.StringIO(text), sep=";", engine="python",
                             quotechar='"', escapechar="\\", on_bad_lines="skip")
        dfs[name] = df

df_registro_classe    = dfs.get("registro_classe.csv")
df_registro_fundo     = dfs.get("registro_fundo.csv")
df_registro_subclasse = dfs.get("registro_subclasse.csv")

df_registro_classe = df_registro_classe.drop_duplicates(subset=["Codigo_CVM"])
df_registro_fundo = df_registro_fundo.drop_duplicates(subset=["Codigo_CVM"])
df_registro_subclasse = df_registro_subclasse.drop_duplicates(subset=["Codigo_CVM"])

#%%


df_registro_fundo["Data_Constituicao"] = pd.to_datetime(
    df_registro_fundo["Data_Constituicao"], errors="coerce"
).dt.strftime("%d/%m/%Y")


df_registro_fundo["Data_Registro"] = pd.to_datetime(
    df_registro_fundo["Data_Registro"], errors="coerce"
)

#%%

def formatar_cpf_cnpj(valor,tipo_valor=""):
    # mantém apenas números
    num = re.sub(r"\D", "", str(valor))

    # Detecta CPF (11 dígitos)
    if tipo_valor == "PF":
       if len(num) > 11:  # remove zeros à esquerda ou excesso
           num = num[-11:]
       if len(num) == 11:
           return f"{num[:3]}.{num[3:6]}.{num[6:9]}-{num[9:]}"
       return valor  # se não tiver 11 dígitos, retorna original
    
    # Corta ou preenche para no máximo 14 dígitos
    if len(num) > 14:
        num = num[:14]
    # Detecta CNPJ (14 dígitos)
    if len(num) < 14:
        num = num.zfill(14)
    if len(num) == 14:
        return f"{num[:2]}.{num[2:5]}.{num[5:8]}/{num[8:12]}-{num[12:]}"
    
    # Se não bate com nenhum formato, retorna como está
    return valor
#%%
# script para gerar docx

from docx import Document

import io, zipfile, re
from docx import Document

def preencher_ficha(data_base: dict) -> io.BytesIO:
    zip_buffer = io.BytesIO()
    with zipfile.ZipFile(zip_buffer, "w", compression=zipfile.ZIP_DEFLATED) as zf:
        for cnpj, data_dict in data_base.items():
            doc = Document(io.BytesIO(TEMPLATE_BYTES))

            for p in doc.paragraphs:
                for key, val in data_dict.items():
                    ph = f"{{{{{key}}}}}"
                    if ph in p.text:
                        p.text = p.text.replace(ph, "" if val is None else str(val))

            for tabela in doc.tables:
                for linha in tabela.rows:
                    for celula in linha.cells:
                        for key, val in data_dict.items():
                            ph = f"{{{{{key}}}}}"
                            if ph in celula.text:
                                celula.text = celula.text.replace(ph, "" if val is None else str(val))

            tmp = io.BytesIO()
            doc.save(tmp); tmp.seek(0)
            cnpj_nome = re.sub(r"\D", "", str(cnpj))
            zf.writestr(f"FICHA_FUNDO_{cnpj_nome}.docx", tmp.getvalue())

    zip_buffer.seek(0)
    return zip_buffer


#%%

# %%
# app.py
import re
import streamlit as st

# --- utilidades ---
def so_digitos(x): 
    return re.sub(r"\D", "", str(x or ""))

def buscar_por_cnpj(cnpj_raw):
    cnpj_digits = so_digitos(cnpj_raw)
    if len(cnpj_digits) != 14:
        return None, "CNPJ inválido: informe 14 dígitos."

    # Normaliza colunas de CNPJ para comparação por dígitos
    # cls = df_registro_fundo.copy()
    # cls["_cnpj_digits"] = cls["CNPJ_Fundo"].astype(str).str.replace(r"\D","", regex=True)

    fdo = df_registro_fundo.copy()
    fdo["_cnpj_digits"] = fdo["CNPJ_Fundo"].astype(str).str.replace(r"\D","", regex=True)

    row_cls = fdo.loc[fdo["_cnpj_digits"] == cnpj_digits].head(1)
    if row_cls.empty:
        return None, "CNPJ não encontrado em registro_classe."
    row_cls = row_cls.iloc[0]

    codigo_cvm = row_cls.get("Codigo_CVM", "")
    rz_social  = row_cls.get("Denominacao_Social", "")
    dt_consti  = row_cls.get("Data_Constituicao", "")
    cnpj_fundo = formatar_cpf_cnpj(cnpj_digits)
    cnpj_adm  = row_cls.get("CNPJ_Administrador", "")
    cnpj_adm = formatar_cpf_cnpj(cnpj_adm)
    adm_nome  = row_cls.get("Administrador", "")
    gestor    = row_cls.get("Gestor", "")
    type_gest = row_cls.get("Tipo_Pessoa_Gestor", "")
    cnpj_gest = row_cls.get("CPF_CNPJ_Gestor", "")
    cnpj_gest = formatar_cpf_cnpj(cnpj_gest,type_gest)




    # Exclusivo (subclasse) via Codigo_CVM
    # sub = df_registro_subclasse
    # row_sub = sub.loc[sub["Codigo_CVM"] == codigo_cvm].head(1) if "Codigo_CVM" in sub.columns else pd.DataFrame()
    # fundo_exclusivo = row_sub.iloc[0].get("Exclusivo", "") if not row_sub.empty else "Não Encontrado"

    dados = {
        "Codigo_CVM": codigo_cvm,
        "Denominacao_Social": rz_social,
        "Data_Constituicao": dt_consti,
        "CNPJ_Fundo": cnpj_fundo,
        "Administrador": adm_nome,
        "CNPJ_Administrador": cnpj_adm,
        "Gestor": gestor,
        "CPF_CNPJ_Gestor": cnpj_gest,
        "Exclusivo": '???',
    }
    return dados, None

def buscar_lote(cnpjs_raw: str):
    # separa por vírgula, ponto e vírgula ou quebra de linha
    itens = re.split(r"[,\n;]+", cnpjs_raw.strip())
    cnpjs = []
    invalidos = []
    for it in itens:
        d = so_digitos(it)
        if not d:
            continue
        if len(d) == 14:
            cnpjs.append(d)
        else:
            invalidos.append(it.strip())

    resultados = []
    erros = []
    vistos = set()
    for d in cnpjs:
        if d in vistos:
            continue
        vistos.add(d)
        dados, err = buscar_por_cnpj(d)
        if err:
            erros.append(f"{formatar_cpf_cnpj(d)}: {err}")
        else:
            resultados.append(dados)

    df_out = pd.DataFrame(resultados)
    # ordem de colunas (se existirem)
    ordem = [
        "Codigo_CVM","Denominacao_Social","Data_Constituicao","CNPJ_Fundo",
        "Administrador","CNPJ_Administrador","Gestor","CPF_CNPJ_Gestor","Exclusivo"
    ]
    df_out = df_out[[c for c in ordem if c in df_out.columns]]
    df_out = df_out.set_index(keys = 'CNPJ_Fundo',drop=False)
    df_out = df_out.convert_dtypes()

    return df_out, invalidos, erros

#%%

# --- UI ---
st.set_page_config(page_title="Consulta CVM por CNPJ — Lote", layout="wide")

# CSS específico de cada item
st.markdown("""
<style>
/* Fundo do app */
.stApp { background: #FFFFFF; }

/* Cabeçalho */
.header-container{
  display:flex;justify-content:space-between;align-items:center;
  padding: 12px 0; border-bottom:1px solid #DDD;
}
.header-title{ font-size:1.25rem; font-weight:700; color:#000; }
.header-logo img{ height:50px; }

/* TextArea (única da página) */
.stTextArea textarea{
  background:d3d3d3 !important; /* cinza escuro */
  color:#FFF !important;          /* texto digitado branco */
  border:2px solid #4682B4 !important;
}
.stTextArea textarea::placeholder{
  color:#FFF !important;          /* placeholder branco */
  opacity:1;
}
/* Label da TextArea (apenas essa) */
label[for="cnpjs_input"]{ color:#000 !important; font-weight:600; }

/* Botão "Consultar" (único) */
.stButton > button{
  background:#0F3B66 !important;  /* azul escuro */
  color:#FFF !important;           /* texto branco */
  border:2px solid #0A2947 !important;
  font-weight:700; border-radius:6px; padding:0.4rem 1rem;
}
.stButton > button:hover{
  background:#0A2947 !important; color:#FFF !important;
}
/* Flash messages (success/warning/error/info) em preto */
.stAlert, .stAlert * { color: #000 !important; }
            
/* Mensagens e tabela: manter legível no fundo branco */
.stAlert, .stDataFrame { color:#000 !important; }
</style>
""", unsafe_allow_html=True)

# Cabeçalho com título e logo
col1, col2 = st.columns([0.2, 0.8])

with col1:
    st.image("logo.png", use_container_width=True)
with col2:
    st.markdown(
    "<h1 style='color:black;'>Consulta de Fundo (CVM) por CNPJ</h1>",
    unsafe_allow_html=True
    )

st.markdown("<hr style='border:1px solid black'>", unsafe_allow_html=True)

# TextArea (com key para estilizar a label específica)

st.markdown(
    "<span style='color:black; font-weight:600;'>Informe 1 ou mais CNPJs (separe por vírgula, ';' ou por linha)</span>",
    unsafe_allow_html=True
)

cnpj_in = st.text_area(
    "\nInforme 1 ou mais CNPJs (separe por vírgula, ';' ou por linha)",
    placeholder="00.000.000/0000-00, 11.111.111/1111-11\n22222222000122",
    height=120,
    key="cnpjs_input",
    label_visibility='collapsed'
)

# Botão azul escuro
import tempfile

# === Consulta ===
if st.button("Consultar"):
    df_out, invalidos, erros = buscar_lote(cnpj_in)
    st.session_state["df_out"] = df_out  # persiste
    st.session_state.pop("zip_bytes", None)  # limpa zip anterior

# === Resultados sempre que existirem ===
if "df_out" in st.session_state and not st.session_state["df_out"].empty:
    df_to_show = st.session_state["df_out"].astype(str)
    st.dataframe(df_to_show.T, use_container_width=True)

    # Gerar ZIP e guardar bytes no estado
    if st.button("Gerar Fichas", key="gerar_zip"):
        df_clean = st.session_state["df_out"].astype(object).where(pd.notnull(st.session_state["df_out"]), None)
        dados_dict = df_clean.T.to_dict()
        zip_buffer = preencher_ficha(dados_dict)   # BytesIO
        zip_buffer.seek(0)
        st.session_state["zip_bytes"] = zip_buffer.getvalue()

# === Download aparece após gerar ===
if "zip_bytes" in st.session_state:
    st.download_button(
        label="📦 Baixar Fichas (.zip)",
        data=st.session_state["zip_bytes"],
        file_name="fichas_fundos.zip",
        mime="application/zip",
        key="download_zip"
    )